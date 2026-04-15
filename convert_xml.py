from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx():
    doc = Document()
    
    # Title Section
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('BPCL')
    run.font.name = 'Arial'
    run.font.size = Pt(52)
    run.font.color.rgb = RGBColor(232, 39, 75)
    run.bold = True
    
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run('Agentic AI Chatbot')
    run.font.name = 'Arial'
    run.font.size = Pt(44)
    run.font.color.rgb = RGBColor(27, 27, 47)
    run.bold = True
    
    plan = doc.add_paragraph('4-Month Implementation Plan')
    plan.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = plan.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(170, 170, 170)
    
    doc.add_page_break()
    
    # Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph('This document defines the end-to-end 4-month implementation roadmap for the BPCL Agentic AI Chatbot, built on the Databricks Data Intelligence Platform. The solution delivers an enterprise-grade, conversational AI interface over 115+ Qlik Sense dashboards with bidirectional data flow, governed by Unity Catalog and secured under a Zero-Trust architecture.')
    doc.add_paragraph('As the sole developer on this engagement, the plan is structured to deliver incremental, demonstrable value at the end of each month, with weekly client check-ins to manage risk, validate scope, and course-correct early.')
    
    doc.save('implementation_plan.docx')

if __name__ == '__main__':
    create_docx()
